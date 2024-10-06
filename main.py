import whisper
from transformers import pipeline
import argparse
from docx import Document
import os

# Set up argparse to accept command-line arguments
parser = argparse.ArgumentParser(description="Transcribe and summarize audio.")
parser.add_argument('audio_file', type=str, help='Path to the audio file (e.g., EarningsCall.wav)')
args = parser.parse_args()

# Load Whisper model and transcribe the audio
model = whisper.load_model("base")
result = model.transcribe(args.audio_file)

# Print transcription
print("Transcription:\n", result["text"])

# Load the BART summarization pipeline
summarizer = pipeline("summarization")

# Sample input text for summarization
input_text = result["text"]

# Generate summary using BART
summary = summarizer(input_text, max_length=500, min_length=200, length_penalty=2.0, num_beams=4, truncation=True)

# Print the generated summary
print("\nGenerated Summary:\n", summary[0]['summary_text'])

# Create a new Word document
doc = Document()

# Add title to the document
doc.add_heading('Transcription and Summary', 0)

# Add transcription to the document
doc.add_heading('Transcription:', level=1)
doc.add_paragraph(result["text"])

# Add summary to the document
doc.add_heading('Summary:', level=1)
doc.add_paragraph(summary[0]['summary_text'])

# Construct output file name
output_file_name = os.path.splitext(args.audio_file)[0] + '.docx'

# Save the document with the constructed name
doc.save(output_file_name)

# Confirmation message
print(f"\nTranscription and summary have been saved to '{output_file_name}'")
